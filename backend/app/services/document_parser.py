import re
from typing import Optional, Dict, Any, List
from abc import ABC, abstractmethod

try:
    import docx
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

try:
    import PyPDF2
    HAS_PYPDF2 = True
except ImportError:
    HAS_PYPDF2 = False

try:
    from PIL import Image
    import pytesseract
    HAS_OCR = True
except ImportError:
    HAS_OCR = False


class DocumentParser(ABC):
    @abstractmethod
    def parse(self, content: bytes) -> Dict[str, Any]:
        pass

    @abstractmethod
    def extract_key_info(self, text: str) -> Dict[str, Any]:
        pass


class TextParser(DocumentParser):
    def parse(self, content: bytes) -> Dict[str, Any]:
        try:
            text = content.decode("utf-8")
        except UnicodeDecodeError:
            try:
                text = content.decode("gbk")
            except UnicodeDecodeError:
                text = content.decode("latin-1", errors="ignore")

        text = self._clean_text(text)
        return {
            "text": text,
            "word_count": len(text.split()),
            "char_count": len(text),
            "format": "txt"
        }

    def extract_key_info(self, text: str) -> Dict[str, Any]:
        keywords = self._extract_keywords(text)
        entities = self._extract_entities(text)
        summary = self._generate_summary(text)

        return {
            "keywords": keywords,
            "entities": entities,
            "summary": summary,
            "language": self._detect_language(text)
        }

    def _clean_text(self, text: str) -> str:
        text = re.sub(r'\s+', ' ', text)
        text = re.sub(r'[\x00-\x08\x0b-\x0c\x0e-\x1f\x7f]', '', text)
        return text.strip()

    def _extract_keywords(self, text: str, top_n: int = 10) -> List[str]:
        stop_words = {
            '的', '了', '在', '是', '我', '有', '和', '就', '不', '人', '都', '一', '一个',
            '上', '也', '很', '到', '说', '要', '去', '你', '会', '着', '没有', '看', '好',
            '自己', '这', '那', '他', '她', '它', '们', '这个', '那个', '什么', '怎么', '为',
            'the', 'a', 'an', 'is', 'are', 'was', 'were', 'be', 'been', 'being', 'have',
            'has', 'had', 'do', 'does', 'did', 'will', 'would', 'could', 'should', 'may',
            'might', 'must', 'can', 'to', 'of', 'in', 'for', 'on', 'with', 'at', 'by',
            'from', 'as', 'or', 'and', 'but', 'if', 'then', 'so', 'because', 'when', 'where'
        }

        words = re.findall(r'[\u4e00-\u9fff]+|[a-zA-Z]+', text)
        word_freq = {}
        for word in words:
            word_lower = word.lower()
            if len(word_lower) >= 2 and word_lower not in stop_words:
                word_freq[word_lower] = word_freq.get(word_lower, 0) + 1

        sorted_words = sorted(word_freq.items(), key=lambda x: x[1], reverse=True)
        return [word for word, _ in sorted_words[:top_n]]

    def _extract_entities(self, text: str) -> Dict[str, List[str]]:
        entities = {
            "organizations": [],
            "locations": [],
            "dates": [],
            "numbers": []
        }

        date_patterns = [
            r'\d{4}[-/年]\d{1,2}[-/月]\d{1,2}[日]?',
            r'\d{1,2}[-/月]\d{1,2}[日]?',
            r'\d{4}年\d{1,2}月',
            r'(?:19|20)\d{2}',
        ]
        for pattern in date_patterns:
            matches = re.findall(pattern, text)
            entities["dates"].extend(matches[:5])

        number_patterns = [
            r'\d+\.?\d*\s*[%℃℃%°]?',
            r'\$\d+(?:,\d{3})*(?:\.\d{2})?',
            r'¥\d+(?:,\d{3})*(?:\.\d{2})?',
            r'\d+(?:,\d{3})+',
        ]
        for pattern in number_patterns:
            matches = re.findall(pattern, text)
            entities["numbers"].extend(matches[:10])

        return entities

    def _generate_summary(self, text: str, max_length: int = 200) -> str:
        sentences = re.split(r'[。！？.!?]+', text)
        sentences = [s.strip() for s in sentences if len(s.strip()) > 10]

        if not sentences:
            return text[:max_length] + "..." if len(text) > max_length else text

        return sentences[0][:max_length] + "..." if len(sentences[0]) > max_length else sentences[0]

    def _detect_language(self, text: str) -> str:
        chinese_count = len(re.findall(r'[\u4e00-\u9fff]', text))
        english_count = len(re.findall(r'[a-zA-Z]', text))

        if chinese_count > english_count:
            return "zh"
        elif english_count > chinese_count:
            return "en"
        else:
            return "mixed"


class DocxParser(DocumentParser):
    def parse(self, content: bytes) -> Dict[str, Any]:
        if not HAS_DOCX:
            return {
                "text": "[DOCX parsing not available - python-docx not installed]",
                "word_count": 0,
                "char_count": 0,
                "format": "docx",
                "error": "Library not available"
            }

        import io
        try:
            doc = docx.Document(io.BytesIO(content))
            paragraphs = []
            for para in doc.paragraphs:
                if para.text.strip():
                    paragraphs.append(para.text)

            text = "\n".join(paragraphs)
            tables_text = []
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text for cell in row.cells)
                    if row_text.strip():
                        tables_text.append(row_text)

            if tables_text:
                text += "\n\n[Tables]\n" + "\n".join(tables_text)

            return {
                "text": text,
                "word_count": len(text.split()),
                "char_count": len(text),
                "format": "docx",
                "paragraph_count": len(paragraphs),
                "table_count": len(doc.tables)
            }
        except Exception as e:
            return {
                "text": f"[Error parsing DOCX: {str(e)}]",
                "word_count": 0,
                "char_count": 0,
                "format": "docx",
                "error": str(e)
            }

    def extract_key_info(self, text: str) -> Dict[str, Any]:
        base_parser = TextParser()
        return base_parser.extract_key_info(text)


class PdfParser(DocumentParser):
    def parse(self, content: bytes) -> Dict[str, Any]:
        if not HAS_PYPDF2:
            return {
                "text": "[PDF parsing not available - PyPDF2 not installed]",
                "word_count": 0,
                "char_count": 0,
                "format": "pdf",
                "error": "Library not available"
            }

        import io
        try:
            from PyPDF2 import PdfReader
            reader = PdfReader(io.BytesIO(content))
            text_parts = []
            page_count = len(reader.pages)

            for page_num, page in enumerate(reader.pages):
                text = page.extract_text()
                if text:
                    text_parts.append(f"[Page {page_num + 1}]\n{text}")

            full_text = "\n\n".join(text_parts)

            return {
                "text": full_text,
                "word_count": len(full_text.split()),
                "char_count": len(full_text),
                "format": "pdf",
                "page_count": page_count
            }
        except Exception as e:
            return {
                "text": f"[Error parsing PDF: {str(e)}]",
                "word_count": 0,
                "char_count": 0,
                "format": "pdf",
                "error": str(e)
            }

    def extract_key_info(self, text: str) -> Dict[str, Any]:
        base_parser = TextParser()
        info = base_parser.extract_key_info(text)

        if "page_count" in text:
            page_match = re.search(r'\[Page \d+\]', text)
            if page_match:
                info["source"] = "PDF Document"
        return info


class ImageParser:
    def parse(self, content: bytes, filename: str = "") -> Dict[str, Any]:
        if not HAS_OCR:
            return {
                "text": "[Image OCR not available - Pillow/Tesseract not installed]",
                "format": "image",
                "error": "OCR library not available"
            }

        try:
            from PIL import Image
            import io
            import pytesseract

            img = Image.open(io.BytesIO(content))

            if img.mode != 'RGB':
                img = img.convert('RGB')

            text = pytesseract.image_to_string(img, lang='chi_sim+eng')

            width, height = img.size
            format_type = filename.rsplit('.', 1)[-1].lower() if '.' in filename else 'unknown'

            return {
                "text": text.strip(),
                "word_count": len(text.split()),
                "char_count": len(text),
                "format": format_type,
                "image_info": {
                    "width": width,
                    "height": height,
                    "mode": img.mode,
                    "format": img.format
                }
            }
        except Exception as e:
            return {
                "text": f"[Error processing image: {str(e)}]",
                "format": "image",
                "error": str(e)
            }

    def extract_key_info(self, text: str) -> Dict[str, Any]:
        base_parser = TextParser()
        return base_parser.extract_key_info(text)


def get_document_parser(file_format: str) -> DocumentParser:
    format_lower = file_format.lower().replace('.', '')
    parser_map = {
        'txt': TextParser,
        'docx': DocxParser,
        'doc': DocxParser,
        'pdf': PdfParser,
    }
    parser_class = parser_map.get(format_lower)
    if parser_class:
        return parser_class()
    return TextParser()
